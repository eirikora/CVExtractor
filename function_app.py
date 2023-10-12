import azure.functions as func
import logging
import os
import tempfile
import json
import base64
import mammoth
import quopri
import re
import html
from bs4 import BeautifulSoup
from urllib.parse import unquote
from docx import Document


app = func.FunctionApp(http_auth_level=func.AuthLevel.ANONYMOUS)

def settKryssCleanup(document_text):
    return_text = ""
    innenfor = False
    last_line = ""
    for line in document_text.split("\n"):
        # print('READ:'+line)
        if innenfor:
            if line.strip() == "":
                innenfor = False
                last_line = ""
            elif line.strip().endswith(':'):
                innenfor = False
                last_line = ""
            elif line.strip() == "Personlige egenskaper":
                innenfor = False
                last_line = ""
            elif line.strip().lower() == "x":
                return_text += last_line + "\n"
                last_line = ""
            else:
                last_line = line
                # print('KILLING:'+line)
        if not innenfor:
            cleanline = line.lower().replace(".","").replace(":","").replace(";","").strip()
            if cleanline.find("sett kryss") > -1: # or cleanline.find("foreslått erfaringsnivå") > -1:
                innenfor = True
            else:
                return_text += line + "\n"
    return return_text

def identifyFooters(document):
    # Count frequency first
    frequencybank = {}
    maxfrequency = 0
    countuniquelines = 0
    for textline in document.split("\n"):
        # Remove EN_SPACE characters
        textline = textline.replace("\u2002", " ")
        # Remove space before end of line
        textline = textline.rstrip()
        if textline != "":
            if textline in frequencybank.keys():
                frequencybank[textline] += 1
                if frequencybank[textline] > maxfrequency:
                    maxfrequency = frequencybank[textline]
            else:
                frequencybank[textline] = 1
                countuniquelines += 1
    frequencyarray = []
    for i in frequencybank.values():
        frequencyarray.append(i)
    freqdistrib = {}
    for item in frequencyarray:
        freqdistrib[item] = freqdistrib.get(item, 0) + 1
    steps = []
    for i in freqdistrib.keys():
        steps.append(i)
    listlength = len(steps)
    footers = []
    if listlength > 0:
        steps.sort(reverse = True)
        accumulatedcount = 0
        for i in range(listlength):
            accumulatedcount += freqdistrib[steps[i]]
            if accumulatedcount > 5 or accumulatedcount > (countuniquelines // 20): # Less than 5% of all lines should be footers
                break
        i -= 1
        if i < 0:
            # no footers
            threshold = maxfrequency + 1
        else:
            threshold = steps[i]    
            for phrase in frequencybank.keys():
                if frequencybank[phrase] >= threshold:
                    footers.append(phrase)
    #print("Threshold frequency set at: "+ str(threshold) + " resulting in " + str(len(footers)) + " footers.")
    return footers

charactermap = { "\u2002":" ",
                "\u201c":"\"",
                "\u201d":"\"",
                "\u2013":"-",
                "\u2014":"-"
                    }
    
def cleanDocument(document_content):
    # Correct encoding errors often seen in HTML content
    document_content = html.unescape(document_content)

    # Remove excessively repeating footers
    footers = [] # Disabled problematic code: identifyFooters(document_content)
    
    # Fjern innholdsfortegnelse
    document_wo_fortegnelse = ""
    inside_innholdsfortegnelse = False
    for textline in document_content.split("\n"):
        cleanline = textline.lower().strip().rstrip(":")
        if cleanline == "innholdsfortegnelse" or cleanline == "innhold":
            # print("FANT INNHOLDSFORTEGNELSE")
            inside_innholdsfortegnelse = True
        else:
            if inside_innholdsfortegnelse:
                if cleanline == "":
                    # print("TOM LINJE")
                    inside_innholdsfortegnelse = False
                elif not cleanline[-1].isdigit():
                    # print("Ikke digit:"+cleanline)
                    inside_innholdsfortegnelse = False
        if not inside_innholdsfortegnelse:
            document_wo_fortegnelse += textline + "\n"
        #else:
        #    print("IGNORING INDEX:" + textline + "\n")
    
    # Traverse document again line by line to remove duplicates and fix characters
    emptylines = 0
    seen_content = False
    result_content = ""
    memorybank = {}
    duplicateCount = 0
    totalDuplicateCount = 0
    buffer = ""
    previousline = ""
    for textline in document_wo_fortegnelse.split("\n"):
        # Remove/replace unicode characters
        for (code, translated) in charactermap.items():
            textline = textline.replace(code, translated)
        # Remove space before end of line
        textline = textline.rstrip()
        # Fix for Bane NOR: Edit Fagområde heading in table if it actually Erfaringsnivå to avoid confusion for ChatGPT
        if (previousline == "Foreslått erfaringsnivå:" and textline == "Fagområde"):
            textline = "" # Skip confusing work Fagområde when after Foreslått erfaringsnivå:
        previousline = textline

        # Remove duplicate lines
        if textline == "" or textline in footers:
            emptylines += 1
        else:
            # Identify and remove duplicates
            seen_content = True
            emptylines = 0
            if textline in memorybank.keys():
                if (textline.endswith(':')): # Insert extra lineshift ahead of headings
                    buffer += "¨"
                buffer += textline + "¨"
                memorybank[textline] += 1
                seen_content = False
                duplicateCount += 1
                # print(str(duplicateCount) + " duplicate " + str(memorybank[textline]) + ":" + textline)
                if duplicateCount >= 3: # After 3 repeating lines start emptying/deleting buffer
                    buffer = ""
            else:
                memorybank[textline] = 1
                totalDuplicateCount += duplicateCount
                duplicateCount = 0
            
        if seen_content and emptylines < 3:
            # Only outputs content once we have seen it (avoid trailing blanks) and only accepts 2 empty lines
            if (textline.endswith(':')): # Insert extra lineshift ahead of headings
                result_content += buffer + "¨" + textline + "¨"
            else:
                result_content += buffer + textline + "¨"
            buffer = ""
    totalDuplicateCount += duplicateCount
    
    if len(footers) > 0:
        result_content += "FOOTERS:¨"
        for line in footers:
            result_content += line + "¨"
    
    result_content = re.sub("^¨+", "", result_content) # Remove unnessesary lineshifts at beginning of document
    result_content = re.sub(":¨+", ":¨", result_content) # Remove multiple lineshifts after colon (:). One is sufficient.
    result_content = re.sub("¨¨+", "¨¨", result_content) # Remove more than two lineshifts in sequence. Two is always sufficient.
    result_content = re.sub("¨", "\n", result_content) # Put back the lineshift character
        
    if totalDuplicateCount > 0:
        print( str(totalDuplicateCount)+ " duplicates removed!")

    # Interpret "Sett kryss" sections correctly
    final_result = settKryssCleanup(result_content)
    if final_result == "\n":
        final_result = ""
    return final_result

def headingToRegex(heading):
    # TARGET REGEX is
    # (^|[^T]¤\\n|[^¤]\\n)(([\.0-9]+\s+)?(\s*[a-zA-Z0-9]\)\s+)?HEADING:?[^\\$]*)
    # where:
    # (^|[^T]¤\\n|[^¤]\\n) = Start of document or start of line, BUT NOT preceeded by T¤ which would indicate inside an existing match
    # ([\.0-9]+\s+)?(\s*[a-zA-Z0-9]\)\s+)? = Optional chapter number (e.g. "5.2") or "b)" or both
    # HEADING:? = Heading text with optional ":" at end
    # [^\\$]* = Some other characters until end of line or document
    # Trim space before and after
    heading = heading.strip()
    if heading.endswith(":"):
        heading = heading.rstrip(":")
    # Escape characters that need escaping
    heading = heading.replace("/", "\/")
    heading = heading.replace("^", "\^")
    heading = heading.replace("[", "\[")
    heading = heading.replace("]", "\]")
    heading = heading.replace("-", "\-")

    # Find generic elements in heading
    heading = re.sub("^\s*[\.0-9]+\s*", "", heading)
    heading = re.sub("^\s*[a-zA-Z0-9]\)\s+", "", heading)
    heading = re.sub("[ \n\t]", "¤SPACE¤", heading)
    heading = heading.replace("*", "¤STAR¤")
    heading = heading.replace(")", "¤CLOSEBRACKET¤")
    heading = heading.replace("(", "¤OPENBRACKET¤")
    heading = heading.replace("?", "¤QUESTION¤")
    heading = re.sub("\\.[0-9]", "¤NUMBER¤", heading)
    heading = re.sub("[0-9]", "¤NUMBER¤", heading)
    heading = heading.replace(".", "\.")
    # Remove duplicates
    heading = re.sub("(¤SPACE¤)+", "¤SPACE¤", heading)
    heading = re.sub("(¤NUMBER¤)+", "¤NUMBER¤", heading)
    # Insert correct Regex
    heading = heading.replace("¤MAYBENUMBER¤", "([\\.0-9]*|[a-zA-Z0-9]?\)?")
    heading = heading.replace("¤SPACE¤", "\\s*")
    heading = heading.replace("¤NUMBER¤", "[\\.0-9]+")
    heading = heading.replace("¤STAR¤", "\\*")
    heading = heading.replace("¤CLOSEBRACKET¤", "\\)")
    heading = heading.replace("¤OPENBRACKET¤", "\\(")
    heading = heading.replace("¤QUESTION¤", "\\?")
    heading = heading.replace("¤ANYTHING¤", "(.*)")
    if ":" not in heading:
        # Add optional end colon for most headings
        heading = heading + ":?"
    matchStartOfLine = "(^|[^T]¤¨|[^¤]¨)"
    matchOptionalChapternum = "([\.0-9]+\s+)?(\s*[a-zA-Z0-9]\)\s+)?"
    matchRestOfLine = "[^¨$]*"
    regex = matchStartOfLine + "(" + matchOptionalChapternum + heading + matchRestOfLine + ")" # Brackets mark Match group 2
    return regex

bucketstore = {}

def writeBlockToBuckets(theblock):
    # Clean up the  text block so that it can be stored
    textHeading = re.sub("\n", "", theblock)
    textHeading = re.sub("¤HEADSTART¤(.+)¤BUCKETINFO¤.*¤HEADEND¤.*$", "\g<1>", textHeading)
    # info "HEADING:" + textHeading
    targetbucketstring = re.sub("\n", "", theblock)
    targetbucketstring = re.sub("^.+¤BUCKETINFO¤(.*)¤HEADEND¤.*$", "\g<1>", targetbucketstring)
    targetbuckets = json.loads(targetbucketstring)
    # info "TARGET BUCKETS:" + targetbuckets
    textBody = re.sub("\n", "¤LF¤", theblock)
    textBody = re.sub("¤HEADSTART¤.*¤BUCKETINFO¤.*¤HEADEND¤", "", textBody)
    textBody = textBody.replace("¤LF¤", "\n")
    # info "BODY:" + textBody
    finalBlock = textHeading + textBody + "\n"
    # Write the last text block to the selected buckets
    #print("DUMP " + finalBlock + " to " + str(targetbuckets))
    for bucket in targetbuckets:
        if finalBlock not in bucketstore[bucket]:
            bucketstore[bucket] += finalBlock

@app.route(route="DocSeparator", auth_level=func.AuthLevel.ANONYMOUS)
def DocSeparator(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('DocSeparator trigger function processed a request.')
    # logging.info('Request is:'+str(req.get_body()))
    debug_block = ""

    try:
        req_body = req.get_json()
    except ValueError:
        return func.HttpResponse(
             "Could not find request body in json format. Please try again!:"+str(req.get_body()),
             status_code=400
        )
    
    if not 'headingmap' in req_body or len(req_body['headingmap']) == 0:
        return func.HttpResponse(
             "Could not find a valid headingmap in the request. Please try again!",
             status_code=400
        )
    
    if not 'texttoseparate' in req_body or len(req_body['texttoseparate']) == 0:
        return func.HttpResponse(
             "Could not find a valid texttoseparate in the request. Please try again!",
             status_code=400
        )
    
    headingmap = req_body.get('headingmap')
    textToSeparate = req_body.get('texttoseparate')

    # Remove all duplicates and mess from document conversion (especially pdf2text which creates duplicates)
    # textToSeparate = cleanDocument(textToSeparate)
    
    # Find all buckets and set them up. Calculate regexmap
    allbuckets = ['Kundebeskrivelse', 'Oppdragsbeskrivelse'] # At least these need to be there for the code to work.
    regexmap = {}
    bucketmap = {}
    headinglist = {}
    headingcount = 0
    for (heading, buckets) in headingmap.items():
        headingcount += 1
        for bucket in buckets:
            if bucket not in allbuckets:
                allbuckets.append(bucket)
        regexforheading = headingToRegex(heading)
        # logging.info(heading + " got regex " + regexforheading)
        regexmap[headingcount] = regexforheading
        bucketmap[headingcount] = buckets
        headinglist[headingcount] = heading
    # print(allbuckets)

    #debug_block += "REGEX USED:\n"
    #for i in regexmap.keys():
    #    debug_block += str(i) + ": " + headinglist[i] + " => " + regexmap[i] + "\n"
    #debug_block += "--- END REGEX LIST ---\n\n"

    for bucket in allbuckets:
        bucketstore[bucket] = ""

    # print(bucketstore)
    #logging.info("REGEXMAP:"+str(regexmap))
    #logging.info("BUCKETMAP:"+str(bucketmap))

    # Capture first 3 lines of doc for Kundebeskrivelse and Oppdragsbeskrivelse (just in case)
    """ 
    linenum = 0
    firstlines = ""
    for textline in textToSeparate.split("\n"):
        if len(textline.strip()) > 0:
            if any(re.findall(r'mvh|hilsen', textline, re.IGNORECASE)):
                break
            linenum += 1
            if linenum <= 3:
                firstlines += textline + "\n"
            else:
                break

    bucketstore['Kundebeskrivelse'] += firstlines + "\n"
    bucketstore['Oppdragsbeskrivelse'] += firstlines + "\n" """

    # Check all triggers against the document. If firing write also previous line to that bucket.
    triggerList = []
    if 'triggerlist' in req_body:
        triggerList = req_body.get('triggerlist')
        previousline = ""
        remaining_lines_to_show = 0
        previous_triggered_buckets = []
        for textline in textToSeparate.split("\n"):
            triggered = False
            triggered_buckets = []
            for trigger in triggerList:
                flags = re.IGNORECASE if ("Ignorer" in trigger["case_sensitivitet"]) else 0
                if any(re.findall(trigger["regex"], textline, flags)):
                    triggered = True
                    for bucket in trigger["til_buckets"]:
                        if bucket not in triggered_buckets:
                            triggered_buckets.append(bucket)
            if triggered:
                debug_block += "TRIGGERED" + str(triggered_buckets) + ":" + textline + "\n" 
                for bucket in triggered_buckets:
                    # insert textline in this bucket
                    bucketstore[bucket] += previousline + textline + "\n"
                    # print(bucket + ": " + previousline + textline)
                    # bucketstore[bucket] += "TRIGGERED:"
                previousline = "" # Avoid repeating if next line also triggers!
                remaining_lines_to_show = 1
            else:
                previousline = textline + "\n"
                if remaining_lines_to_show > 0:
                    for bucket in previous_triggered_buckets:
                        # insert textline in this bucket
                        bucketstore[bucket] += textline + "\n"
                    remaining_lines_to_show -= 1
            previous_triggered_buckets = triggered_buckets.copy()

    textToSeparate = re.sub("\n","¨", textToSeparate) # Insert special character for EOL
    # logging.info("INITIAL:"+textToSeparate)

    # Match all the heading regex against the textfile to identify the location of the headings
    for headingnum in regexmap.keys():
        buckets = json.dumps(bucketmap[headingnum])
        textToSeparate = re.sub(regexmap[headingnum], "\g<1>¤HEADSTART¤" + "\g<2>" + "¤BUCKETINFO¤" + buckets + "¤HEADEND¤", textToSeparate, flags=re.IGNORECASE) # The heading is now match group 2. Match group is what we tested BEFORE heading.
        # Debug version
        # textToSeparate = re.sub(regexmap[headingnum], "\g<1>¤HEADSTART¤" + "\g<2>" + "//" + headinglist[headingnum] + "//" + str(headingnum) + "¤BUCKETINFO¤" + buckets + "¤HEADEND¤¨", textToSeparate, flags=re.IGNORECASE) # The heading is now match group 2. Match group is what we tested BEFORE heading.
        #print("\n")
        #print(headingnum)
        #print(headinglist[headingnum])
        #print(regexmap[headingnum])
        
    #logging.info("MATCHED:"+textToSeparate)
    debug_block += "MATCHES IN DOCUMENT:\n"
    for debugline in textToSeparate.split("¨"):
        if "¤" in debugline and "\.\.\.\.\.\.\.\.\.\." not in debugline: # The "...." indicates an index which is not relevant
            debug_block += debugline + "\n"
    debug_block += "--- END MATCHES IN DOCUMENT ---\n\n"

    # Separate the text file: Identify all text blocks and write them to the right buckets
    insideBlock = False
    blockContent = ""
    for textline in textToSeparate.split("¨"):
        if insideBlock:
            if "¤HEADSTART¤" in textline and "\.\.\.\.\.\.\.\.\.\." not in textline: # The "...." indicates an index which is not relevant
                # Skriv blokken til buckets som er ønsket
                writeBlockToBuckets(blockContent)
                # Tøm tekstblokken og start på nytt med denne overskriftslinjen
                blockContent = textline + "\n"
            else:
                # Just another line of text to keep...
                blockContent = blockContent + textline + "\n"
        else:
            if "¤HEADSTART¤" in textline and "\.\.\.\.\.\.\.\.\.\." not in textline: # The "...." indicates an index which is not relevant
                # Vi starter en ny tekstblokk
                insideBlock = True
                blockContent = textline + "\n"
                
    if insideBlock:
        # Skriv også siste blokken til buckets som er ønsket
        writeBlockToBuckets(blockContent)

    #Clean up the text before returning it (as Zoho is so poor at processing a lot of text).
    for bucket in bucketstore.keys():
        bucketstore[bucket] = cleanDocument(bucketstore[bucket])

    bucketstore["DEBUGINFO"] = re.sub("¤HEADSTART¤([^¤]+)¤BUCKETINFO¤([^¤]*)¤HEADEND¤", "\g<1> ==> \g<2>", debug_block)

    return func.HttpResponse(json.dumps(bucketstore), mimetype="application/json", status_code=200)

def decode_mime_string(s):
    if '=' in s:
        start = s.find('=')
        prefix = s[:start]
        #print(prefix)
        mimestring = s[start:]
        #print(mimestring)
        mime_elements = mimestring.split('=_=')
        #print(mime_elements)
        decoded_text = ''
        for submime in mime_elements:
            charset, encoded_text = submime.strip('=_').strip('_=').split('_Q_', 1)
            #decoded_text += quopri.decodestring(encoded_text.replace('_', '=')).decode(charset).replace('=', ' ')
            decoded_text += quopri.decodestring(encoded_text).decode(charset)
        return prefix + decoded_text
    else:
        return s


@app.route(route="MimeDecode", auth_level=func.AuthLevel.ANONYMOUS)
def MimeDecode(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('MimeDecode trigger function processed a request.')

    mime_string = req.params.get('mime_string')
    if not mime_string:
        try:
            req_body = req.get_json()
        except ValueError:
            mime_string = str(req.get_body(), 'utf-8')
        else:
            mime_string = req_body.get('mime_string')

    if mime_string:
        decoded_string = decode_mime_string(mime_string) # Now also handles prefixed mime-strings
        return func.HttpResponse(decoded_string, mimetype="text/plain;charset=UTF-8", status_code=200)
    else:
        return func.HttpResponse(
             "Please pass a mime_string on the query string or in the request body",
             status_code=400
        )

def allowed_file(filename):
    logging.info("Checking "+ filename)
    ALLOWED_EXTENSIONS = {'doc', 'docx'}
    return '.' in filename and filename.rsplit('.', 1)[1].lower() in ALLOWED_EXTENSIONS

def extract_header_footer(docx_path):
    # Load the document using python-docx
    doc = Document(docx_path)

    # Extract header text
    header_text = []
    for section in doc.sections:
        for header in section.header.paragraphs:
            header_text.append(header.text)

    # Extract footer text
    footer_text = []
    for section in doc.sections:
        for footer in section.footer.paragraphs:
            footer_text.append(footer.text)

    return "\n".join(header_text + footer_text)

@app.route(route="Word2Text", auth_level=func.AuthLevel.ANONYMOUS)
def Word2Text(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('Word2Text trigger function processed a request.')

    file_content = b''
    file_name = ''
    if 'content' not in req.files:
        try:
            logging.info('Getting the body...')
            req_body = req.get_body()
            json_data = json.loads(req_body, strict=False)
            encoded_content = json_data["parameters"]["body"]["$content"]
            # Decode the content
            file_content = base64.b64decode(encoded_content)
            file_name = "powerapp.docx"
        except ValueError:
            logging.info('ERROR: No file part in the request.')
            #file_content = req.get_body()
            return func.HttpResponse('ERROR: No file part in the request', mimetype="text/plain;charset=UTF-8", status_code=400)
        #return 'No file part in the request', 400
        #return func.HttpResponse(str(req_body[0:100]), mimetype="text/plain;charset=UTF-8", status_code=400)
    else:
        file = req.files['content']
        file_content = file.read()
        file_name = file.filename

    # Expect a Word filename to attempt convert
    if allowed_file(file_name):
        logging.info('Converting file ' + file_name)
        # Write the binary data to a temp Word file
        with tempfile.NamedTemporaryFile(mode='w+b', delete=False, suffix=".docx") as f:
            f.write(file_content)
            temp_filename = f.name
        
        # Get the header and footer from the file
        headerfooter = extract_header_footer(temp_filename)

        # Read and convert temp Word file into HTML for analysis
        with open(temp_filename, "rb") as docx_file:
            try:
                result = mammoth.convert_to_html(docx_file)
                html = result.value
            except Exception as e:
                html = "<p>Word2Text ERROR: Document was not a Microsoft Word document in .docx format (Zip file) that could be analyzed.</p>\n"
                # hex_representation = file_content[0:200].hex()
                # hex_representation_spaced = ' '.join(hex_representation[i:i+2] for i in range(0, len(hex_representation), 2))
                # html = html + "<p>" + hex_representation_spaced + "</p>\n"
                # html = html + "<p>" + str(req_body)[2:40] + "</p>\n"
        # Delete temp file
        try:
            os.remove(temp_filename)
        except OSError:
            pass
        
        # Create a BeautifulSoup object
        soup = BeautifulSoup(html, 'html.parser')

        plain_text = soup.get_text(separator='\n')
        plain_text = "HEADER_FOOTER:\n" + headerfooter + '\nFRONT_PAGE:\n' + plain_text
        plain_text = re.sub('\t',' ',plain_text)
        # Return the resultset
        return func.HttpResponse(plain_text, mimetype="text/plain;charset=UTF-8", status_code=200)
    else:
        logging.info('ERROR: Did not find a Word document to convert in request body!')
        return func.HttpResponse(
             "ERROR: Could not find a file in doc/docx format that we can convert. File received was: " + file_name,
             status_code=400
        )


@app.route(route="ExtractCVData", auth_level=func.AuthLevel.ANONYMOUS)
def ExtractCVData(req: func.HttpRequest) -> func.HttpResponse:
    logging.info('ExtractCVData trigger function processed a request.')

    # Get the CV body word data
    word_data=b''
    try:
        logging.info('Getting the body...')
        req_body = req.get_body()
    except ValueError:
        logging.info('Body has a value error.')
        word_data = b''
        pass
    else:
        logging.info('Getting body.')
        word_data = req_body
    logging.info('Got a document to analyze of length '+str(len(word_data)))

    #json_data='nada'
    #json_data = json.loads(word_data, strict=False)
    decoded_data = b''

    # Try to decode the data in case it's base64-encoded
    try:
        # Parse JSON and extract content
        logging.info(f"Trying to decode...")
        #json_data = word_data
        json_data = json.loads(word_data, strict=False)
        encoded_content = json_data["parameters"]["body"]["$content"]  #json_data.get('$content')
        logging.info('Got the encoded string starting with ' + encoded_content[0:10])
        # Decode base64-encoded content
        decoded_data = base64.b64decode(encoded_content)
        logging.info("Data was base64-encoded. Successfully decoded.") 
    except Exception as e:
        logging.info(f"Data was not encoded. Treating it as binary data.")
        decoded_data = word_data

    # Expect at least 20 bytes
    if (len(decoded_data) > 20):
        # Write the binary data to a temp Word file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
            f.write(decoded_data)
            temp_filename = f.name

        # Read and convert temp Word file into HTML for analysis
        with open(temp_filename, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value

        # Delete temp file
        os.remove(temp_filename)

        standard_headings = ['Nøkkelkompetanse', 'Kompetanse', 'key competency', 'kvalifikasjoner', 'qualifications',
            'Prosjekterfaring', 'Erfaring', 'experience', 'Utdanning', 'Utdannelse', 'education', 
            'Kurs', 'courses', 'Kurs/Sertifiseringer', 'courses/certifications', 'Annet', 'other',
            'Spesialkunnskap', 'Relevant spesialkunnskap', 'relevant special knowledge', 
            'Arbeidsgivere', 'employers', 'Språk', 'languages']
        
        # Create a BeautifulSoup object
        soup = BeautifulSoup(html, 'html.parser')

        # Find all <p> tags
        paragraphs = soup.find_all('p')

        # Iterate over each <p> tag to identify all headings and mark them with <h1>
        for paragraph in paragraphs:
            if isinstance(paragraph.string, str):
                # Check if the paragraph can be a recognizable heading
                for standard_heading in standard_headings:
                    if re.match("^"+standard_heading.lower(), paragraph.string.lower().strip()):
                        # Create a new <h1> tag and replace the paragraph with it
                        bold_tag = soup.new_tag("h1")
                        bold_tag.string = paragraph.string.strip()
                        paragraph.replace_with(bold_tag)
                        break
        
        #Find all h1 headings again
        allheadings=[]
        headings = soup.find_all('h1')
        for tag in headings:
            heading = tag.get_text().strip()
            allheadings.append(heading)

        # Heading Translation map
        translationdict = {'Nøkkelkompetanse':'key competency', 
            'Kompetanse':'key competency',
            'Kvalifikasjoner':'qualifications',
            'Prosjekterfaring':'experience', 
            'Erfaring':'experience', 
            'Utdanning':'education', 
            'Utdannelse':'education', 
            'Kurs':'courses/certifications', 
            'Kurs/Sertifiseringer':'courses/certifications', 
            'Annet':'other',
            'Spesialkunnskap':'relevant special knowledge', 
            'Relevant spesialkunnskap':'relevant special knowledge', 
            'Arbeidsgivere':'employers', 
            'Språk':'languages'}
        language = 'English'

        # Collect all text per heading/section
        resultset = {} #"docname":docname, "consultantname":consultantname }
        #logging.info("All headings found:" + repr(allheadings))
        for tag in headings:
            heading = tag.get_text().replace(':','').strip()
            headingkey = heading.lower()
            if heading in translationdict.keys():
                headingkey = translationdict[heading].lower()
                language = 'Norwegian'

            # print("\nExtracting SECTION: "+heading)
            sectioncontent = ""
            fathertag = tag #.parent # Need to go one level up (p-tag)
            # Collect text until next heading in CV
            for sibling in fathertag.next_siblings:
                if sibling.get_text().strip() in allheadings:
                    # End of current section
                    resultset[headingkey] = sectioncontent
                    sectioncontent = ""
                    break
                else: # Add one more line to current section
                    for textline in sibling.stripped_strings:
                        sectioncontent += textline+'\n'
            if (len(sectioncontent) > 0):
                resultset[headingkey] = sectioncontent

        resultset['cv language'] = language

        # Return the resultset
        return func.HttpResponse(json.dumps(resultset), mimetype="application/json", status_code=200)# Expect at least 20 bytes
    if (len(decoded_data) > 20):
        # Write the binary data to a temp Word file
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as f:
            f.write(decoded_data)
            temp_filename = f.name

        # Read and convert temp Word file into HTML for analysis
        with open(temp_filename, "rb") as docx_file:
            result = mammoth.convert_to_html(docx_file)
            html = result.value

        # Delete temp file
        os.remove(temp_filename)

        standard_headings = ['Nøkkelkompetanse', 'Kompetanse', 'key competency', 'kvalifikasjoner', 'qualifications',
            'Prosjekterfaring', 'Erfaring', 'experience', 'Utdanning', 'Utdannelse', 'education', 
            'Kurs', 'courses', 'Kurs/Sertifiseringer', 'courses/certifications', 'Annet', 'other',
            'Spesialkunnskap', 'Relevant spesialkunnskap', 'relevant special knowledge', 
            'Arbeidsgivere', 'employers', 'Språk', 'languages']
        
        # Create a BeautifulSoup object
        soup = BeautifulSoup(html, 'html.parser')

        # Find all <p> tags
        paragraphs = soup.find_all('p')

        # Iterate over each <p> tag to identify all headings and mark them with <h1>
        for paragraph in paragraphs:
            if isinstance(paragraph.string, str):
                # Check if the paragraph can be a recognizable heading
                for standard_heading in standard_headings:
                    if re.match("^"+standard_heading.lower(), paragraph.string.lower().strip()):
                        # Create a new <h1> tag and replace the paragraph with it
                        bold_tag = soup.new_tag("h1")
                        bold_tag.string = paragraph.string.strip()
                        paragraph.replace_with(bold_tag)
                        break
        
        #Find all h1 headings again
        allheadings=[]
        headings = soup.find_all('h1')
        for tag in headings:
            heading = tag.get_text().strip()
            allheadings.append(heading)

        # Heading Translation map
        translationdict = {'Nøkkelkompetanse':'key competency', 
            'Kompetanse':'key competency',
            'Kvalifikasjoner':'qualifications',
            'Prosjekterfaring':'experience', 
            'Erfaring':'experience', 
            'Utdanning':'education', 
            'Utdannelse':'education', 
            'Kurs':'courses/certifications', 
            'Kurs/Sertifiseringer':'courses/certifications', 
            'Annet':'other',
            'Spesialkunnskap':'relevant special knowledge', 
            'Relevant spesialkunnskap':'relevant special knowledge', 
            'Arbeidsgivere':'employers', 
            'Språk':'languages'}
        language = 'English'

        # Collect all text per heading/section
        resultset = {} #"docname":docname, "consultantname":consultantname }
        #logging.info("All headings found:" + repr(allheadings))
        for tag in headings:
            heading = tag.get_text().replace(':','').strip()
            headingkey = heading.lower()
            if heading in translationdict.keys():
                headingkey = translationdict[heading].lower()
                language = 'Norwegian'

            # print("\nExtracting SECTION: "+heading)
            sectioncontent = ""
            fathertag = tag #.parent # Need to go one level up (p-tag)
            # Collect text until next heading in CV
            for sibling in fathertag.next_siblings:
                if sibling.get_text().strip() in allheadings:
                    # End of current section
                    resultset[headingkey] = sectioncontent
                    sectioncontent = ""
                    break
                else: # Add one more line to current section
                    for textline in sibling.stripped_strings:
                        sectioncontent += textline+'\n'
            if (len(sectioncontent) > 0):
                resultset[headingkey] = sectioncontent

        resultset['cv language'] = language

        # Return the resultset
        return func.HttpResponse(json.dumps(resultset), mimetype="application/json", status_code=200)
    else:
        logging.info('Did not find a CV in request body!')
        return func.HttpResponse(
             "This HTTP triggered function executed successfully, but lacked cv in docx format to analyze.",
             status_code=400
        )
